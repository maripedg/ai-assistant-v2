#!/usr/bin/env python3
"""
DOCX Sanitizer (local)

- Preserves images/diagrams and table structure by editing text *runs* only
- Skips headings (Heading 1/2/3...) to avoid section-number false positives
- Sanitizes:
  [email], [url], [key], [ip], [ipv6], [customer], [location]
- Generates a text log with counts per placeholder

Usage:
  python sanitize_docx.py "input.docx" "output_sanitized.docx" "sanitization_log.txt"
"""

import re
import sys
from collections import defaultdict
from docx import Document


# -------------------------
# Configurable patterns
# -------------------------
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
URL_RE   = re.compile(r"https?://[^\s)>\"]+")
# Long hex-like tokens (adjust if you need more/less strict)
KEY_RE   = re.compile(r"\b[A-Fa-f0-9]{32,}\b")

# IPv6: requires at least 2 colons and hex chars
IPV6_RE  = re.compile(r"\b(?=[0-9A-Fa-f:]{6,}\b)(?:[0-9A-Fa-f]{1,4}:){2,7}[0-9A-Fa-f]{1,4}\b")
# Candidate IPv4
IPV4_RE  = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")

# Optional: customer / location patterns (customize as needed)
# You can add multiple customer names or cities here.
CUSTOMER_TERMS = ["Rakuten"]
LOCATION_TERMS = ["Sanda", "Totsuka", "Okayama"]

CUSTOMER_RE = re.compile(r"\b(" + "|".join(map(re.escape, CUSTOMER_TERMS)) + r")\b", re.IGNORECASE) if CUSTOMER_TERMS else None
LOCATION_RE = re.compile(r"\b(" + "|".join(map(re.escape, LOCATION_TERMS)) + r")\b", re.IGNORECASE) if LOCATION_TERMS else None

# Context keywords: if an IPv4 looks like section numbering (e.g., 2.1.4.1),
# only replace if IP-ish terms appear in the same paragraph.
IP_CONTEXT_KEYWORDS = re.compile(r"\b(ip|ipv4|ipv6|address|addr|subnet|gateway|gw|dns)\b", re.IGNORECASE)


def ipv4_is_real_ip(ip_str: str, context: str) -> bool:
    """
    Heuristic:
    - Validate octets 0..255.
    - Always treat private/infra ranges as IPs: 10/172.16-31/192.168/127.*
    - Avoid section numbering false positives by requiring context when first octet <= 25.
    """
    parts = ip_str.split(".")
    try:
        octets = [int(p) for p in parts]
    except ValueError:
        return False
    if any(o < 0 or o > 255 for o in octets):
        return False

    first = octets[0]
    # Common internal ranges
    if first == 10 or first == 127 or first == 192 or (first == 172 and 16 <= octets[1] <= 31):
        return True

    # Likely a section number (2.1.4.1, 6.2.3.1...) unless context says otherwise
    if first <= 25:
        return bool(IP_CONTEXT_KEYWORDS.search(context))

    return True


def sanitize_text(text: str, context: str, log: defaultdict) -> str:
    def _sub(regex, repl, key):
        nonlocal text
        matches = regex.findall(text)
        if matches:
            log[key] += len(matches)
            text = regex.sub(repl, text)

    _sub(EMAIL_RE, "[email]", "[email]")
    _sub(URL_RE,   "[url]",   "[url]")
    _sub(KEY_RE,   "[key]",   "[key]")

    if CUSTOMER_RE:
        _sub(CUSTOMER_RE, "[customer]", "[customer]")
    if LOCATION_RE:
        _sub(LOCATION_RE, "[location]", "[location]")

    # IPv6
    def ipv6_repl(_m):
        log["[ipv6]"] += 1
        return "[ipv6]"
    text = IPV6_RE.sub(ipv6_repl, text)

    # IPv4
    def ipv4_repl(m):
        ip = m.group(0)
        if ipv4_is_real_ip(ip, context):
            log["[ip]"] += 1
            return "[ip]"
        return ip
    text = IPV4_RE.sub(ipv4_repl, text)

    return text


def is_heading(paragraph) -> bool:
    try:
        return bool(paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading"))
    except Exception:
        return False


def sanitize_paragraph(paragraph, log: defaultdict):
    # Skip headings to avoid breaking numbering (optional but recommended)
    if is_heading(paragraph):
        return

    # Build full paragraph context for IPv4 heuristic
    context = "".join(run.text for run in paragraph.runs)

    # IMPORTANT: modify runs only (preserves images/objects and table structure)
    for run in paragraph.runs:
        if run.text:
            run.text = sanitize_text(run.text, context, log)


def sanitize_docx(input_path: str, output_path: str, log_path: str):
    doc = Document(input_path)
    log = defaultdict(int)

    # Body paragraphs
    for p in doc.paragraphs:
        sanitize_paragraph(p, log)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sanitize_paragraph(p, log)

    # Headers/footers: by default we DO sanitize them.
    # If you want to skip them entirely, comment this block.
    for section in doc.sections:
        for p in section.header.paragraphs:
            sanitize_paragraph(p, log)
        for p in section.footer.paragraphs:
            sanitize_paragraph(p, log)

    doc.save(output_path)

    with open(log_path, "w", encoding="utf-8") as f:
        for k in sorted(log.keys()):
            f.write(f"{k}: {log[k]}\n")


def main():
    if len(sys.argv) != 4:
        print("Usage: python sanitize_docx.py <input.docx> <output.docx> <log.txt>")
        sys.exit(2)

    input_path, output_path, log_path = sys.argv[1], sys.argv[2], sys.argv[3]
    sanitize_docx(input_path, output_path, log_path)
    print(f"Sanitized DOCX written to: {output_path}")
    print(f"Log written to: {log_path}")


if __name__ == "__main__":
    main()
