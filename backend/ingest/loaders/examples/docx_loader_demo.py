from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from backend.ingest.loaders.docx_loader import load


def build_sample_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("Title Block", level=1)
    doc.add_paragraph("Intro paragraph under title.")
    doc.add_heading("1 Main Section", level=1)
    doc.add_heading("1.1 Overview", level=2)
    doc.add_paragraph("Overview content line 1.")
    doc.add_paragraph("Overview content line 2.")
    doc.add_heading("2 Second Section", level=1)
    doc.add_paragraph("Second section content.")
    doc.save(path)


def main():
    with TemporaryDirectory() as td:
        docx_path = Path(td) / "demo.docx"
        build_sample_doc(docx_path)
        items = load(str(docx_path))
        for idx, item in enumerate(items, start=1):
            print(f"--- Item {idx} ---")
            print("heading_path:", item["metadata"].get("heading_path"))
            print("section_heading:", item["metadata"].get("section_heading"))
            print("heading_level_of_section:", item["metadata"].get("heading_level_of_section"))
            print("numbering_prefix_of_section:", item["metadata"].get("numbering_prefix_of_section"))
            print("text:\n", item["text"])


if __name__ == "__main__":
    main()
