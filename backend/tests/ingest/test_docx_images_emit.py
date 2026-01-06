import base64
from io import BytesIO
from pathlib import Path

from docx import Document

from backend.ingest.loaders import docx_loader
from backend.ingest.loaders.chunking.toc_section_docx_chunker import chunk_docx_toc_sections


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAoAAAAKCAQAAACEN8SrAAAAJElEQVR42mP8z/D/PwMDAwMjI+P///8ZGBgY/j8GhgYGBgYGAF6VBi3G1tF/AAAAAElFTkSuQmCC"
)


def _build_doc_with_image(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Before image")
    img_stream = BytesIO(PNG_BYTES)
    doc.add_picture(img_stream)
    doc.add_paragraph("After image")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return path


def test_docx_images_emit_pipeline(tmp_path, monkeypatch):
    doc_path = _build_doc_with_image(tmp_path)
    assets_dir = tmp_path / "assets"
    monkeypatch.setenv("DOCX_EXTRACT_IMAGES", "0")
    monkeypatch.setenv("DOCX_INLINE_FIGURE_PLACEHOLDERS", "1")
    monkeypatch.setenv("DOCX_FIGURE_CHUNKS", "1")
    monkeypatch.setenv("RAG_ASSETS_DIR", str(assets_dir))

    items = docx_loader.load(str(doc_path))
    image_items = [it for it in items if (it.get("metadata") or {}).get("block_type") == "image"]
    assert image_items, "image blocks should be emitted when flags enable figure handling"

    chunks = chunk_docx_toc_sections(items, cfg={"effective_max_tokens": 512}, source_meta={"doc_id": "sample"})
    placeholders = [c for c in chunks if "[FIGURE:" in c.get("text", "")]
    figure_chunks = [c for c in chunks if (c.get("metadata") or {}).get("chunk_type") == "figure"]
    assert placeholders, "placeholders should be inserted when inline flag is on"
    assert figure_chunks, "figure chunks should be emitted when flag is on"
    for fig in figure_chunks:
        assert fig.get("metadata", {}).get("parent_chunk_id")
