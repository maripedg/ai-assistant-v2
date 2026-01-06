import base64
from io import BytesIO
from pathlib import Path

from docx import Document

from backend.ingest.loaders import docx_loader


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAoAAAAKCAQAAACEN8SrAAAAJElEQVR42mP8z/D/PwMDAwMjI+P///8ZGBgY/j8GhgYGBgYGAF6VBi3G1tF/AAAAAElFTkSuQmCC"
)


def _build_doc_with_image(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Before image")
    img_stream = BytesIO(PNG_BYTES)
    doc.add_picture(img_stream)
    doc.add_paragraph("After image")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return path


def test_docx_image_extraction_and_block_order(tmp_path, monkeypatch):
    doc_path = _build_doc_with_image(tmp_path)
    assets_dir = tmp_path / "assets"
    monkeypatch.setenv("DOCX_EXTRACT_IMAGES", "1")
    monkeypatch.setenv("RAG_ASSETS_DIR", str(assets_dir))

    items = docx_loader.load(str(doc_path))

    written = list((assets_dir / "sample").glob("img_*.png"))
    assert written, "image files should be written to assets dir"
    figure_id = "sample_img_001"
    expected_ref = f"sample/{written[0].name}"
    figure_items = [it for it in items if (it.get("metadata") or {}).get("block_type") == "image"]
    assert figure_items, "image item should be emitted"
    assert figure_items[0]["text"] == figure_id
    image_ref = figure_items[0]["metadata"].get("image_ref")
    assert image_ref == expected_ref
    assert not Path(image_ref).is_absolute()
    assert (assets_dir / image_ref).exists()

    texts = [it["text"] for it in items]
    before_idx = texts.index(next(t for t in texts if "Before image" in t))
    image_idx = texts.index(figure_id)
    after_idx = texts.index(next(t for t in texts if "After image" in t))
    assert before_idx < image_idx < after_idx


def test_docx_images_skipped_when_flag_off(tmp_path, monkeypatch):
    doc_path = _build_doc_with_image(tmp_path)
    assets_dir = tmp_path / "assets_off"
    monkeypatch.setenv("DOCX_EXTRACT_IMAGES", "0")
    monkeypatch.setenv("RAG_ASSETS_DIR", str(assets_dir))

    items = docx_loader.load(str(doc_path))

    assert not (assets_dir / "sample").exists()
    assert all((it.get("metadata") or {}).get("block_type") != "image" for it in items)
